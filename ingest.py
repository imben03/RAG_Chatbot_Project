# ingest.py - Document Ingestion and Chunking Pipeline

# Key functions
#   1. Scan the docs/ folder for PDF, DOCX, HTML files
#   2. Extract raw text from each document
#   3. Split text into overlapping chunks for better retrieval
#   4. Generate vector embeddings via Google Gemini API
#   5. Store chunks + embeddings in ChromaDB
#   6. Track file hashes in index_log.json to skip unchanged files
#   7. Auto-detect versioned documents (by year) and keep only the latest

import os
import re
import time
import logging
import hashlib
import json
from pathlib import Path
from datetime import datetime
import PyPDF2
import docx
from bs4 import BeautifulSoup
import chromadb
from google import genai
from dotenv import load_dotenv

load_dotenv()

# logging setup
# Writes detailed logs to chatbot.log; only WARNING+ shown in console
logging.basicConfig(
    filename='chatbot.log',
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
console = logging.StreamHandler()
console.setLevel(logging.WARNING)
logging.getLogger().addHandler(console)

# Configuration constants
CHUNK_SIZE    = 500   # Target number of words per chunk
CHUNK_OVERLAP = 50    # Words shared between consecutive chunks for context continuity
COLLECTION    = 'mum_policy_docs'   # ChromaDB collection name
EMBED_DELAY   = 0.5   # Seconds to wait between embedding calls (rate limit protection)
DB_PATH       = './chroma_store'    # ChromaDB persistence directory
INDEX_LOG     = './index_log.json'  # Tracks file hashes and indexing metadata

#Google Gemini API client setup
# Uses v1beta API version — required for gemini-embedding-001 model access
client_genai = genai.Client(
    api_key=os.getenv('GOOGLE_API_KEY'),
    http_options={'api_version': 'v1beta'}
)

# Index log helpers

def load_index_log():
    """
    Load the persistent file index log from disk.
    Returns an empty dict if the log file does not exist yet.
    The log maps filename -> {hash, chunks, indexed_at, file_type, doc_title}.
    """
    if os.path.exists(INDEX_LOG):
        with open(INDEX_LOG, 'r') as f:
            return json.load(f)
    return {}

def save_index_log(log):
    """
    Persist the updated index log back to disk after each file is processed.
    Called incrementally so progress is not lost if ingestion is interrupted.
    """
    with open(INDEX_LOG, 'w') as f:
        json.dump(log, f, indent=2)

def file_hash(path):
    """
    Compute MD5 hash of a file's raw bytes.
    Used to detect whether a file has changed since it was last indexed,
    so unchanged files can be skipped without re-embedding.
    """
    h = hashlib.md5()
    with open(path, 'rb') as f:
        h.update(f.read())
    return h.hexdigest()

# ──────────────────────────────────────────────────────────────────────
# Version Detection — Auto-replace outdated documents
# ──────────────────────────────────────────────────────────────────────

# Matches 4-digit years (2000–2099) anywhere in a filename
YEAR_PATTERN = re.compile(r'(20\d{2})')

def extract_year(filename):
    """
    Extract the most recent 4-digit year (2000-2099) from a filename.
    Returns the year as an integer, or None if no year is found.

    Examples:
        'Student_Handbook_2025.pdf'       -> 2025
        'Fees_Policy_2023-2024.pdf'       -> 2024  (takes the latest year)
        'General_Regulations.pdf'         -> None   (no year present)
    """
    matches = YEAR_PATTERN.findall(filename)
    if matches:
        return max(int(y) for y in matches)
    return None

def get_base_name(filename):
    """
    Strip year patterns and file extension to get a normalised base name.
    This is used to group different versions of the same document together.

    Steps:
    1. Remove the file extension
    2. Remove all 4-digit year patterns (2000-2099)
    3. Collapse separators (hyphens, underscores, spaces) into a single underscore
    4. Strip leading/trailing underscores
    5. Convert to lowercase for case-insensitive matching

    Examples:
        'Student_Handbook_2025.pdf'       -> 'student_handbook'
        'Student_Handbook_2023.pdf'       -> 'student_handbook'  (same group)
        'Fees-Policy-2023-2024.docx'      -> 'fees_policy'
        'General_Regulations.pdf'         -> 'general_regulations'
    """
    stem = Path(filename).stem
    stripped = YEAR_PATTERN.sub('', stem)
    normalised = re.sub(r'[-_\s]+', '_', stripped).strip('_').lower()
    return normalised

def filter_latest_versions(files):
    """
    Given a list of file paths, group them by base name and keep only the
    file with the highest year for each group. Files without a year in their
    name are always kept (they are not considered versioned).

    Returns:
        keep   — list of Path objects to ingest
        remove — dict mapping base_name -> list of outdated Path objects
                 (these will have their old chunks removed from ChromaDB)
    """
    # Separate files into versioned (has a year) and unversioned (no year)
    versioned  = {}   # base_name -> [(year, path), ...]
    unversioned = []

    for fp in files:
        year = extract_year(fp.name)
        if year is not None:
            base = get_base_name(fp.name)
            versioned.setdefault(base, []).append((year, fp))
        else:
            unversioned.append(fp)

    keep   = list(unversioned)  # Always keep non-versioned files
    remove = {}                 # base_name -> [outdated paths]

    for base, entries in versioned.items():
        # Sort by year descending — first entry is the latest version
        entries.sort(key=lambda e: e[0], reverse=True)
        latest_year, latest_fp = entries[0]
        keep.append(latest_fp)

        # Everything else is outdated
        outdated = [fp for yr, fp in entries[1:]]
        if outdated:
            remove[base] = outdated
            print(f'  Version check: keeping {latest_fp.name} (year {latest_year})')
            for old_fp in outdated:
                old_year = extract_year(old_fp.name)
                print(f'    Superseded: {old_fp.name} (year {old_year}) — will be removed')

    return keep, remove


# Embedding helper with rate limit handling

def get_embedding(text, retries=3):
    """
    Generate a vector embedding for a text string using Gemini embedding model.

    Retries up to 3 times on HTTP 429 (rate limit) errors with exponential
    back-off (15s, 30s, 45s). All other errors are raised immediately.

    Returns a list of floats representing the embedding vector.
    """
    for attempt in range(retries):
        try:
            result = client_genai.models.embed_content(
                model='gemini-embedding-001',
                contents=text
            )
            # Extract the flat values list from the first embedding object
            return result.embeddings[0].values
        except Exception as e:
            if '429' in str(e) and attempt < retries - 1:
                # Rate limited — wait and retry with increasing delay
                wait = (attempt + 1) * 15
                print(f'    Rate limited. Waiting {wait}s...')
                time.sleep(wait)
            else:
                logging.error(f'Embedding error: {e}')
                raise
    return []

# Text extraction functions for different file types

def extract_pdf(path):
    """
    Extract plain text from a PDF file using PyPDF2.

    Each page is prefixed with a [Page N] marker so chunk metadata can
    later reference approximate page locations. Null bytes and bullet
    point artifacts common in PDFs are cleaned before returning.
    """
    text = []
    try:
        with open(path, 'rb') as fh:
            reader = PyPDF2.PdfReader(fh)
            for page_num, page in enumerate(reader.pages):
                extracted = page.extract_text()
                if extracted and extracted.strip():
                    # Remove null bytes (\x00) and common PDF bullet artifacts
                    cleaned = extracted.replace('\x00', '').replace('\uf0b7', '-')
                    text.append(f'[Page {page_num + 1}]\n{cleaned}')
    except Exception as e:
        logging.warning(f'PDF extraction error for {path}: {e}')
    return '\n'.join(text)

def extract_docx(path):
    """
    Extract plain text from a Word (.docx) document.

    Processes both paragraph text and table cell content.
    Table rows are joined with ' | ' separators to preserve structure
    as readable text rather than losing the tabular relationship.
    """
    try:
        doc = docx.Document(path)
        paragraphs = []

        # Extract paragraph text, skipping empty paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                paragraphs.append(p.text.strip())

        # Extract table data — each row becomes a pipe-delimited string
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)

        return '\n'.join(paragraphs)
    except Exception as e:
        logging.warning(f'DOCX extraction error for {path}: {e}')
        return ''

def extract_html(path):
    """
    Extract clean readable text from an HTML file using BeautifulSoup.

    Strips script, style, nav, footer, and header tags before extracting
    text to avoid indexing boilerplate UI chrome rather than policy content.
    """
    try:
        with open(path, 'r', encoding='utf-8') as fh:
            soup = BeautifulSoup(fh, 'html.parser')
        # Remove non-content tags before extracting text
        for tag in soup(['script', 'style', 'nav', 'footer', 'header']):
            tag.decompose()
        return soup.get_text(separator='\n')
    except Exception as e:
        logging.warning(f'HTML extraction error for {path}: {e}')
        return ''

# Maps file extensions to their corresponding extractor functions
EXTRACTORS = {
    '.pdf':  extract_pdf,
    '.docx': extract_docx,
    '.html': extract_html,
    '.htm':  extract_html,
}

# Chunking function to split text into overlapping word-based chunks
def chunk_text(text, size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    """
    Split a document's text into overlapping word-based chunks.

    Strategy:
    - Paragraphs (newline-separated) are accumulated until the target
      word count is reached, then flushed as a chunk.
    - The last `overlap` words of each chunk are carried over into the
      next chunk so that sentences spanning chunk boundaries are not lost.
    - Paragraphs longer than `size` words are split directly into
      sub-chunks without accumulation.
    - Chunks shorter than 50 characters are discarded as noise.

    Returns a list of chunk strings ready for embedding.
    """
    paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
    chunks, current, current_len = [], [], 0

    for para in paragraphs:
        tokens = para.split()
        if not tokens:
            continue

        # Handle paragraphs longer than the target chunk size
        if len(tokens) > size:
            # Flush any accumulated words first
            if current:
                chunk = ' '.join(current)
                if len(chunk) > 50:
                    chunks.append(chunk)
                current, current_len = [], 0
            # Slide a window over the long paragraph
            for i in range(0, len(tokens), size - overlap):
                sub = tokens[i:i + size]
                sub_chunk = ' '.join(sub)
                if len(sub_chunk) > 50:
                    chunks.append(sub_chunk)
            continue

        # Flush current accumulation if adding this paragraph would exceed size
        if current_len + len(tokens) > size and current:
            chunk = ' '.join(current)
            if len(chunk) > 50:
                chunks.append(chunk)
            # Carry over the last `overlap` words into the next chunk
            overlap_tokens = ' '.join(current).split()[-overlap:]
            current = overlap_tokens + tokens
            current_len = len(current)
        else:
            # Keep accumulating words into the current chunk
            current.extend(tokens)
            current_len += len(tokens)

    # Flush whatever remains in the accumulator
    if current:
        chunk = ' '.join(current)
        if len(chunk) > 50:
            chunks.append(chunk)

    return chunks

# ChromaDB helper functions for managing existing document IDs and removing stale chunks

def get_existing_ids(collection):
    """
    Fetch all document IDs currently stored in the ChromaDB collection.
    Used to skip chunks that have already been indexed, preventing duplicates.
    Returns an empty set if the collection cannot be queried.
    """
    try:
        result = collection.get()
        return set(result['ids'])
    except Exception:
        return set()

def remove_document_chunks(collection, source_name):
    """
    Delete all chunks in the collection that belong to a specific source file.
    Called when a file's hash has changed, so stale chunks are removed before
    the updated version is re-indexed.
    """
    try:
        result = collection.get(include=['metadatas'])
        # Find IDs whose metadata 'source' field matches the file being removed
        ids_to_delete = [
            result['ids'][i]
            for i, m in enumerate(result['metadatas'])
            if m.get('source') == source_name
        ]
        if ids_to_delete:
            collection.delete(ids=ids_to_delete)
            print(f'  Removed {len(ids_to_delete)} old chunks for {source_name}')
    except Exception as e:
        logging.error(f'Error removing chunks for {source_name}: {e}')

# Main ingestion function that orchestrates the full pipeline

def ingest_directory(doc_dir='docs'):
    """
    Main ingestion entry point. Orchestrates the full pipeline:

    1. Connect to ChromaDB and load the existing index log
    2. Scan the docs/ folder for supported file types
    3. Auto-detect versioned documents and keep only the latest version
       - Outdated versions have their chunks removed from ChromaDB
    4. For each file to keep:
       a. Compute its MD5 hash
       b. Skip if hash matches the log (unchanged file)
       c. Remove old chunks if the file has changed
       d. Extract text using the appropriate extractor
       e. Split text into overlapping chunks
       f. Embed each chunk via Gemini API (with rate-limit protection)
       g. Store chunks, embeddings, and metadata in ChromaDB
       h. Update the index log entry for this file
    5. Print a summary of what was indexed, skipped, or errored
    """
    # Connect to (or create) the persistent ChromaDB store
    chroma     = chromadb.PersistentClient(path=DB_PATH)
    collection = chroma.get_or_create_collection(
        COLLECTION,
        metadata={'hnsw:space': 'cosine'}  # Cosine similarity for semantic search
    )

    # Snapshot current IDs so we can detect duplicates during this run
    existing_ids = get_existing_ids(collection)
    index_log    = load_index_log()

    # Recursively find all files with supported extensions
    all_files = [fp for fp in Path(doc_dir).rglob('*') if fp.suffix.lower() in EXTRACTORS]

    if not all_files:
        print('No supported documents found in docs/ folder.')
        print('Supported formats: PDF, DOCX, HTML, HTM')
        return

    # ── Version Detection ─────────────────────────────────────────────
    # Group files by base name, keep only the latest year for each group.
    # Outdated versions are removed from ChromaDB and the index log.
    print(f'Found {len(all_files)} document(s) — checking for version duplicates...\n')

    files, outdated_map = filter_latest_versions(all_files)

    # Remove outdated versions from the database
    superseded = 0
    for base, old_paths in outdated_map.items():
        for old_fp in old_paths:
            if old_fp.name in index_log or old_fp.name in {
                m.get('source') for m in collection.get(include=['metadatas'])['metadatas']
            }:
                print(f'  Auto-removing outdated: {old_fp.name}')
                remove_document_chunks(collection, old_fp.name)
                existing_ids = get_existing_ids(collection)  # Refresh after deletion
                if old_fp.name in index_log:
                    del index_log[old_fp.name]
                    save_index_log(index_log)
                superseded += 1

    if superseded > 0:
        print(f'\n  Removed {superseded} outdated document version(s) from the database.\n')
    # ──────────────────────────────────────────────────────────────────

    print(f'Proceeding with {len(files)} document(s) to process...')
    print(f'Database path : {DB_PATH}')
    print(f'Already logged: {len(index_log)} document(s)\n')

    # Counters for the final summary
    total_chunks = 0
    skipped      = 0
    updated      = 0
    errors       = 0

    for i, fp in enumerate(files, 1):
        print(f'[{i}/{len(files)}] {fp.name}')
        fhash = file_hash(str(fp))

        #Check if the file has been indexed before and if it has changed since last time
        if fp.name in index_log:
            if index_log[fp.name]['hash'] == fhash:
                # File is identical to what was last indexed — skip entirely
                print(f'  Unchanged since last index — skipping\n')
                skipped += 1
                continue
            else:
                # File has been modified — remove stale chunks before re-indexing
                print(f'  File has changed — re-indexing...')
                remove_document_chunks(collection, fp.name)
                existing_ids = get_existing_ids(collection)  # Refresh after deletion
                updated += 1

        # Text extraction based on file extension
        ext  = fp.suffix.lower()
        text = EXTRACTORS[ext](str(fp))  # Dispatch to the correct extractor

        if not text.strip():
            print(f'  WARNING: No text extracted — skipping\n')
            logging.warning(f'No text extracted from {fp.name}')
            skipped += 1
            continue

        #chunking
        chunks = chunk_text(text)

        if not chunks:
            print(f'  WARNING: No chunks created — skipping\n')
            skipped += 1
            continue

        # Build IDs and filter out any already in the database
        all_ids   = [f'{fp.stem}_{j}' for j in range(len(chunks))]
        new_pairs = [(cid, ch) for cid, ch in zip(all_ids, chunks) if cid not in existing_ids]

        if not new_pairs:
            # Every chunk ID already exists — nothing new to add
            print(f'  All chunks already indexed — skipping\n')
            skipped += 1
            continue
  
        ids_to_add  = [p[0] for p in new_pairs]
        docs_to_add = [p[1] for p in new_pairs]

        print(f'  Embedding {len(docs_to_add)} chunks...')

        try:
            # Embedding with rate limit handling 
            embeddings = []
            for idx, chunk in enumerate(docs_to_add):
                emb = get_embedding(chunk)
                embeddings.append(emb)
                time.sleep(EMBED_DELAY)  # Throttle to avoid rate limit errors
                if (idx + 1) % 10 == 0:
                    print(f'    {idx + 1}/{len(docs_to_add)} embedded...')

            # Metadata for each chunk — includes source filename, chunk index, file type, and a human-readable title
            # Convert filename to a human-readable title (e.g. my-policy_doc -> My Policy Doc)
            doc_title = fp.stem.replace('-', ' ').replace('_', ' ').title()
            metas = [
                {
                    'source':       fp.name,          # Original filename for filtering
                    'chunk_index':  j,                # Position of this chunk within the doc
                    'file_type':    ext,              # .pdf / .docx / .html
                    'doc_title':    doc_title,        # Human-readable title shown in the UI
                    'indexed_at':   datetime.now().isoformat(),
                    'total_chunks': len(docs_to_add), # Total chunks for this document
                }
                for j in range(len(docs_to_add))
            ]

            # ChromaDB add: store the new chunks, their embeddings, and metadata
            collection.add(
                documents=docs_to_add,
                ids=ids_to_add,
                metadatas=metas,
                embeddings=embeddings
            )

            total_chunks += len(docs_to_add)

            # Update the index log for this file with its current hash and metadata
            # Save after every file so partial progress survives interruptions
            index_log[fp.name] = {
                'hash':       fhash,
                'chunks':     len(docs_to_add),
                'indexed_at': datetime.now().isoformat(),
                'file_type':  ext,
                'doc_title':  doc_title,
            }
            save_index_log(index_log)

            print(f'  Done! ({len(docs_to_add)} chunks indexed)\n')
            logging.info(f'Indexed {len(docs_to_add)} chunks from {fp.name}')

        except Exception as e:
            print(f'  ERROR: {e}\n')
            logging.error(f'Failed to index {fp.name}: {e}')
            errors += 1

    # Summary of the ingestion run
    print('=' * 50)
    print(f'Ingestion complete!')
    print(f'New chunks indexed : {total_chunks}')
    print(f'Files skipped      : {skipped}')
    print(f'Files re-indexed   : {updated}')
    print(f'Versions superseded: {superseded}')
    print(f'Errors             : {errors}')
    print(f'Total in database  : {collection.count()}')
    print('=' * 50)


if __name__ == '__main__':
    ingest_directory()